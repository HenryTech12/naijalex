import pdfplumber
import docx
import pytesseract
from typing import Tuple
from pdf2image import convert_from_path
from PIL import Image
from app.services.ocr_service import perform_ocr
from app.utils.text_cleaner import clean_text

def pdf_ocr_fallback(file_path: str) -> str:
    """Convert PDF pages to images and run tesseract OCR."""
    try:
        images = convert_from_path(file_path, dpi=200)
        texts = []
        for img in images:
            text = pytesseract.image_to_string(img, lang='eng', config='--psm 6')
            texts.append(text)
        return "\n\n".join(texts)
    except Exception as e:
        return f"[OCR failed: {str(e)}]"

def ingest_document(file_path: str, file_type: str) -> Tuple[str, bool]:
    """
    Ingest PDF, Image, or DOCX and return clean text + OCR flag.
    """
    raw_text = ""
    ocr_used = False
    
    if file_type == "application/pdf":
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            total_chars = 0
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages_text.append(text)
                total_chars += len(text)
            
            # Fallback to OCR if text extraction is poor
            avg_chars = total_chars / len(pdf.pages) if pdf.pages else 0
            if avg_chars < 100:
                raw_text = pdf_ocr_fallback(file_path)
                ocr_used = True
            else:
                raw_text = "\n".join(pages_text)
                
    elif file_type.startswith("image/"):
        with open(file_path, "rb") as f:
            raw_text = perform_ocr(f.read())
        ocr_used = True
        
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        raw_text = "\n".join(full_text)
        
    else:
        # Plain text
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw_text = f.read()
            
    return clean_text(raw_text), ocr_used
